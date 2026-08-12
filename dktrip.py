import streamlit as st
import pandas as pd
import os
import base64
import re
import unicodedata
from collections import defaultdict
from streamlit_gsheets import GSheetsConnection
import docx
from io import BytesIO

# ================= CẤU HÌNH TRANG & GIAO DIỆN =================
st.set_page_config(page_title="KẾT QUẢ ĐĂNG KÝ TRIP", page_icon="🌿", layout="centered")

url = "https://docs.google.com/spreadsheets/d/1xwALxKMnTwXeP8Vy3BVJTuD0JpF981_mv_b4AjIaigA/edit"

def get_image_base64(path):
    if os.path.exists(path):
        with open(path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode()
    return ""

cover_base64 = get_image_base64("web cover.jpg")
qr_base64 = get_image_base64("TTCK.jpg")
thongtin_base64 = get_image_base64("THÔNG TIN TRIP POST.jpg")

def parse_money(amount):
    if pd.isna(amount): return 0
    if isinstance(amount, bool): return 0 
    try:
        if isinstance(amount, (int, float)): return int(amount)
        s = str(amount).strip()
        if s.endswith('.0'): s = s[:-2]
        s = re.sub(r'[^\d]', '', s)
        return int(s) if s else 0
    except:
        return 0

def format_vnd(amount):
    val = parse_money(amount)
    return f"{val:,}".replace(',', '.') + " VNĐ" if val > 0 else "0 VNĐ"

def normalize_text(text):
    if pd.isna(text): return ""
    return unicodedata.normalize('NFC', str(text)).strip().lower()

# Lấy Group logic (trước dấu ":")
def get_prefix(text):
    if pd.isna(text): return ""
    txt = str(text).strip()
    return txt.split(':')[0].strip() if ':' in txt else txt

# Lấy Tên hiển thị (trước dấu "-")
def get_display_name(text):
    if pd.isna(text): return ""
    txt = str(text).strip()
    return txt.split('-')[0].strip() if '-' in txt else txt

# --- Hàm tạo file Word danh sách ---
def generate_word(df, dot_name):
    doc = docx.Document()
    doc.add_heading(f"DANH SÁCH KHÁCH CHỐT ĐƠN - {dot_name.upper()}", 0)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Họ và Tên'
    hdr_cells[1].text = 'Năm sinh'

    for _, r in df.iterrows():
        for ng in str(r['DS người']).split(','):
            parts = ng.split('-')
            t = parts[0].strip() if len(parts)>0 else ""
            ns = parts[1].strip() if len(parts)>1 else ""
            row_cells = table.add_row().cells
            row_cells[0].text = t
            row_cells[1].text = ns

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

st.markdown("""
    <style>
    .stApp { background-color: #FFFFFF; color: #333333; }
    h1, h2, h3, h4, .stTabs [data-baseweb="tab"] p { color: #2C5E1A !important; font-weight: bold; }
    
    h1.main-title { text-align: center; color: #2C5E1A !important; font-weight: 900; font-size: clamp(22px, 5vw, 36px); word-break: keep-all; line-height: 1.4; margin-bottom: 30px; }
    
    .stButton>button { background-color: #D4AF37; color: #1A3C0F; font-weight: bold; border-radius: 8px; border: none; width: 100%; padding: 10px; }
    .stButton>button:hover { background-color: #2C5E1A; color: #FFFFFF; border: none; }
    .stTextInput>div>div>input { background-color: #F8F9FA; color: #333333; border: 1px solid #2C5E1A; border-radius: 5px; font-size: 16px; font-weight: bold;}
    
    .section-title { background: linear-gradient(90deg, #2C5E1A 0%, #D4AF37 100%); color: white; padding: 12px 15px; border-radius: 8px 8px 0 0; font-size: 16px; font-weight: bold; margin-top: 25px; text-transform: uppercase; }
    .custom-table { width: 100%; border-collapse: separate; border-spacing: 0; margin-bottom: 20px; border: 1px solid #E0E6ED; border-top: none; border-radius: 0 0 8px 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.05); overflow: hidden; }
    .custom-table thead tr { background-color: #1A3C0F; } 
    .custom-table th { color: white; padding: 12px 14px; text-align: center; font-size: 15px; border: none; }
    .custom-table td { padding: 14px; border-bottom: 1px solid #EEEEEE; border-right: 1px solid #EEEEEE; text-align: center; font-weight: bold; color: #2C5E1A; background-color: #FFFFFF;}
    
    .cancel-alert { background-color: #F8D7DA; color: #721C24; padding: 20px; border-radius: 8px; border: 1px solid #F5C6CB; margin-bottom: 20px; text-align: center; line-height: 1.6; }
    
    .info-card { background: linear-gradient(135deg, #F9FBE7 0%, #FFFDE7 100%); padding: 20px; border-left: 6px solid #2C5E1A; border-radius: 8px; margin-bottom: 20px; box-shadow: 0 2px 5px rgba(0,0,0,0.05); }
    .info-card p { margin-bottom: 8px; font-size: 16px; }
    .info-card strong { color: #2C5E1A; }
    .info-card span { color: #D4AF37; font-weight: bold; font-size: 17px;}
    
    .metric-card { height: 100%; min-height: 180px; display: flex; flex-direction: column; justify-content: center; background-color: #FFFFFF; border: 1px solid #E0E6ED; border-top: 5px solid #2C5E1A; padding: 15px; border-radius: 8px; margin-bottom: 15px; box-shadow: 0 2px 4px rgba(0,0,0,0.05); text-align: center;}
    .metric-title { font-size: 16px; color: #2C5E1A; font-weight: bold; text-transform: uppercase; margin-bottom: 5px;}
    .metric-value { font-size: 28px; color: #D4AF37; font-weight: 900;}
    
    .total-row td { background-color: #F4C430 !important; color: #1A3C0F !important; font-size: 16px; font-weight: 900 !important;}
    .profit-row td { background-color: #E8F5E9 !important; color: #1B5E20 !important; font-size: 18px; font-weight: 900 !important;}
    </style>
""", unsafe_allow_html=True)

conn = st.connection("gsheets", type=GSheetsConnection)

@st.cache_data(ttl=5)
def load_data():
    df_data = conn.read(spreadsheet=url, worksheet="Data App")
    df_config = conn.read(spreadsheet=url, worksheet="Thông số")
    df_data['SDT'] = df_data['SDT'].astype(str).str.replace(r'\.0$', '', regex=True).str.strip()
    
    if 'Link Zalo' not in df_config.columns:
        df_config['Link Zalo'] = ""
        
    return df_data, df_config

try:
    df_data, df_config = load_data()
except Exception as e:
    st.error(f"Lỗi kết nối dữ liệu: {e}")
    st.stop()

col_dot = 'Đợt ĐK' if 'Đợt ĐK' in df_data.columns else 'ĐĂNG KÝ ĐỢT TRIP'
df_data['Original_Index'] = df_data.index

try:
    time_ck = int(df_config.iloc[0]['Thời gian CK']) if not pd.isna(df_config.iloc[0]['Thời gian CK']) else 0
    time_check = int(df_config.iloc[0]['Thời gian Check']) if not pd.isna(df_config.iloc[0]['Thời gian Check']) else 0
    wait_time = time_ck + time_check
except:
    wait_time = 15

# ================= PHẦN 1: USER - KẾT QUẢ ĐĂNG KÝ =================
if cover_base64:
    st.markdown(f"<div style='text-align:center; margin-bottom: 25px;'><img src='data:image/jpeg;base64,{cover_base64}' style='width:100%; border-radius:12px; box-shadow: 0 4px 8px rgba(0,0,0,0.1);'></div>", unsafe_allow_html=True)
else:
    st.markdown("<h1 class='main-title'>KẾT QUẢ ĐĂNG KÝ<br>TRIP ĐÊM HUYỀN BÍ</h1>", unsafe_allow_html=True)

st.markdown("<p style='color: #2C5E1A; font-weight: 900; font-size: 16px; text-transform: uppercase; margin-bottom: 5px;'>NHẬP SỐ ĐIỆN THOẠI CỦA BẠN (ĐÃ DÙNG ĐĂNG KÝ):</p>", unsafe_allow_html=True)
phone_input = st.text_input("SDT", label_visibility="collapsed", placeholder="Ví dụ: 0901234567")

if st.button("TRA CỨU KẾT QUẢ 🚀"):
    if phone_input:
        clean_input = phone_input.strip().lstrip('0')
        df_data['Phone_Compare'] = df_data['SDT'].str.lstrip('0')
        matched_rows = df_data[df_data['Phone_Compare'] == clean_input]
        
        if not matched_rows.empty:
            classified_rows = []
            first_name = "Bạn"
            
            for idx, row in matched_rows.iterrows():
                dot_dk_full = str(row[col_dot]).strip()
                dot_dk_prefix = get_prefix(dot_dk_full)
                sdt = row['SDT']
                sdt_display = f"0{str(sdt).lstrip('0')}"
                status = str(row['Trạng thái CK']).strip().upper()
                sl_dk = int(row['SL']) if pd.notna(row['SL']) else 1
                
                ds_nguoi = str(row['DS người']).split(',')
                if first_name == "Bạn" and ds_nguoi and ds_nguoi[0] != 'nan':
                    parts = ds_nguoi[0].split('-')
                    if parts[0].strip(): first_name = parts[0].strip()
                
                df_dot = df_data[df_data[col_dot].apply(get_prefix).apply(normalize_text) == normalize_text(dot_dk_prefix)].sort_values(by='Timestamp').copy()
                df_dot['CumSum_SL'] = pd.to_numeric(df_dot['SL'], errors='coerce').fillna(0).cumsum()
                
                limit_mask = df_config['Nội dung option'].apply(get_prefix).apply(normalize_text) == normalize_text(dot_dk_prefix)
                limit_series = df_config.loc[limit_mask, 'SL giới hạn']
                
                try:
                    max_slot = int(float(limit_series.values[0])) if not limit_series.empty and pd.notna(limit_series.values[0]) else 30
                except:
                    max_slot = 30
                
                user_cumsum = df_dot.loc[df_dot['Original_Index'] == idx, 'CumSum_SL'].values[0]
                
                cat = status
                lo_type = 0
                if user_cumsum > max_slot:
                    cat = "LO_SLOT"
                    lo_type = 1 if (user_cumsum - sl_dk) < max_slot else 2
                        
                classified_rows.append({
                    'row': row,
                    'dot_dk_full': dot_dk_full,
                    'dot_dk_prefix': dot_dk_prefix,
                    'status': status,
                    'sl_dk': sl_dk,
                    'ds_nguoi': [n for n in ds_nguoi if n.strip() and n != 'nan'],
                    'cat': cat,
                    'lo_type': lo_type,
                    'sdt_display': sdt_display
                })
            
            groups = defaultdict(list)
            for r in classified_rows:
                groups[r['cat']].append(r)
                
            for cat, items in groups.items():
                if cat == "LO_SLOT":
                    for item in items:
                        if item['lo_type'] == 1:
                            st.error(f"Xin lỗi {first_name}, do số lượng đăng ký tại {item['dot_dk_prefix']} ({item['sl_dk']} người) đã vượt quá số lượng slot còn lại. Anh chị liên hệ Zalo Quang - 0902800318 để được hỗ trợ giải quyết nha!")
                        else:
                            st.error(f"Thành thật xin lỗi {first_name} 😭 Do ở cùng một thời điểm có quá nhiều người cùng gửi form nên đăng ký của bạn tại {item['dot_dk_prefix']} được ghi nhận khi đợt trip đã kín chỗ. Anh chị liên hệ Zalo Quang - 0902800318 để được hỗ trợ giải quyết nha!")
                
                elif cat == "HỦY SLOT":
                    dots = " & ".join(list(set([i['dot_dk_prefix'] for i in items])))
                    st.markdown(f"<div class='cancel-alert'><strong>Xin lỗi, đăng ký của {first_name} tại {dots} đã bị HỦY vì quá thời hạn thanh toán mất rồi.</strong></div>", unsafe_allow_html=True)
                
                elif cat in ["ĐANG CẬP NHẬT", "CHỐT ĐƠN THÀNH CÔNG"]:
                    total_sl = sum(i['sl_dk'] for i in items)
                    total_tien_val = sum(parse_money(i['row']['Tổng tiền']) for i in items)
                    
                    # --- Đã update: Cắt phần sau dấu gạch ngang "-" ---
                    dots_full = " <br> ".join(list(set([get_display_name(i['dot_dk_full']) for i in items])))
                    
                    sdt_disp = items[0]['sdt_display']
                    all_ds = []
                    for i in items: all_ds.extend(i['ds_nguoi'])
                    
                    latest_row = items[-1]['row']
                    han_chot = latest_row['Hạn chót CK']
                    tg_con = latest_row['Thời gian còn lại']
                    
                    zalo_link = ""
                    if 'Link Zalo' in df_config.columns:
                        zalo_mask = df_config['Nội dung option'].apply(get_prefix).apply(normalize_text) == normalize_text(items[0]['dot_dk_prefix'])
                        zalo_series = df_config.loc[zalo_mask, 'Link Zalo']
                        if not zalo_series.empty and pd.notna(zalo_series.values[0]):
                            raw_zalo = str(zalo_series.values[0]).strip()
                            if raw_zalo != "" and raw_zalo.lower() not in ["nan", "none"]:
                                zalo_link = raw_zalo
                    
                    info_card_html = f"""
                    <div class='info-card'>
                        <p><strong>Đợt Đăng ký:</strong> <br><span style='font-size: 15px;'>{dots_full}</span></p>
                        <p><strong>SĐT người đại diện:</strong> <span>{sdt_disp}</span></p>
                        <p style='margin-bottom:0;'><strong>Tổng số lượng:</strong> <span>{total_sl} người</span></p>
                    </div>
                    """
                    
                    if cat == "ĐANG CẬP NHẬT":
                        st.success(f"🎉 Chúc mừng {first_name} đã đăng ký thành công, thông tin đăng ký của bạn như sau:")
                        st.markdown(info_card_html, unsafe_allow_html=True)
                        
                        table_html = "<table class='custom-table'><thead><tr><th>Họ và Tên</th><th>Năm sinh</th></tr></thead><tbody>"
                        for ng in all_ds:
                            parts = ng.split('-')
                            t = parts[0].strip() if len(parts)>0 else ""
                            ns = parts[1].strip() if len(parts)>1 else ""
                            table_html += f"<tr><td>{t}</td><td>{ns}</td></tr>"
                        table_html += "</tbody></table>"
                        st.markdown(table_html, unsafe_allow_html=True)
                        
                        noidung_ck = f"TRIP - {sdt_disp}"
                        st.markdown("<div class='section-title'>THÔNG TIN THANH TOÁN</div>", unsafe_allow_html=True)
                        
                        with st.container(border=True):
                            col_info, col_qr = st.columns([1.5, 1])
                            with col_info:
                                st.markdown(f"<div style='margin-bottom: 12px; font-size:15px;'><strong>💰 Tổng số tiền:</strong> <span style='color:#D4AF37; font-weight:bold; font-size:16px;'>{format_vnd(total_tien_val)}</span></div>", unsafe_allow_html=True)
                                st.markdown(f"<div style='margin-bottom: 12px; font-size:15px;'><strong>⏳ Hạn chót CK:</strong> <span style='color:#D4AF37; font-weight:bold; font-size:16px;'>{han_chot}</span></div>", unsafe_allow_html=True)
                                st.markdown(f"<div style='margin-bottom: 18px; font-size:15px;'><strong>⏱️ Thời gian còn lại:</strong> <span style='color:#D4AF37; font-weight:bold; font-size:16px;'>{tg_con}</span></div>", unsafe_allow_html=True)
                                
                                st.markdown("<div style='font-weight: bold; color: #2C5E1A; margin-bottom: 5px; font-size:15px;'>THÔNG TIN CHUYỂN KHOẢN:</div>", unsafe_allow_html=True)
                                st.markdown("<div style='font-size: 15px; margin-bottom: 5px; color: #2C5E1A;'>Tô Văn Quang - Vietcombank<br>STK:</div>", unsafe_allow_html=True)
                                st.code("0251001799405")
                                
                                st.markdown("<div style='font-weight: bold; color: #2C5E1A; margin-bottom: 5px; margin-top: 15px; font-size:15px;'>NỘI DUNG CHUYỂN KHOẢN:</div>", unsafe_allow_html=True)
                                st.code(noidung_ck)
                                st.markdown("<div style='text-align: center; color: red; font-weight: bold; font-size: 15px; margin-top: 5px;'>⚠️ Vui lòng chuyển đúng nội dung</div>", unsafe_allow_html=True)
                            
                            with col_qr:
                                if qr_base64:
                                    st.markdown(f"<div style='text-align:center; padding-top:20px;'><img src='data:image/jpeg;base64,{qr_base64}' style='width:100%; max-width:220px; border-radius:8px; box-shadow: 0 4px 8px rgba(0,0,0,0.1);'></div>", unsafe_allow_html=True)
                                else:
                                    st.write("Đang cập nhật QR...")
                        st.warning(f"🔄 Sau khi chuyển khoản xong, chậm nhất {wait_time} phút sau kết quả nhận chuyển khoản sẽ được cập nhật. Bạn làm mới (refresh) trang để xem kết quả nhé.")
                    
                    elif cat == "CHỐT ĐƠN THÀNH CÔNG":
                        st.success(f"🎉 Chúc mừng {first_name} đã CHỐT ĐƠN THÀNH CÔNG, hoàn tất việc đăng ký.")
                        st.markdown(info_card_html, unsafe_allow_html=True)
                        
                        table_html = "<table class='custom-table'><thead><tr><th>Họ và Tên</th><th>Năm sinh</th></tr></thead><tbody>"
                        for ng in all_ds:
                            parts = ng.split('-')
                            t = parts[0].strip() if len(parts)>0 else ""
                            ns = parts[1].strip() if len(parts)>1 else ""
                            table_html += f"<tr><td>{t}</td><td>{ns}</td></tr>"
                        table_html += "</tbody></table>"
                        st.markdown(table_html, unsafe_allow_html=True)
                        
                        st.markdown(f"""
                        <div style='background-color: #E8F5E9; padding: 15px; border-radius: 8px; border: 1px solid #A5D6A7; margin-bottom: 20px;'>
                            <p style='margin: 0; color: #2E7D32; font-size: 15px;'>✅ BTC đã nhận được thanh toán:</p>
                            <p style='margin: 5px 0 0 0; font-size: 18px; color: #1B5E20;'><strong>Tổng số tiền đã nhận:</strong> <span style='color: #D4AF37;'>{format_vnd(total_tien_val)}</span></p>
                        </div>
                        """, unsafe_allow_html=True)

                        # --- Đã update: Khung Zalo rực rỡ, nằm ngay sau khung thanh toán ---
                        if zalo_link:
                            st.markdown(f"""
                            <div style='margin-bottom: 25px; font-size: 16px; color: #1A3C0F; background: linear-gradient(135deg, #F4D03F 0%, #F1C40F 100%); padding: 22px; border-radius: 10px; border: 3px dashed #E67E22; font-weight: 900; text-align: center; text-transform: uppercase; box-shadow: 0 6px 12px rgba(0,0,0,0.15);'>
                                🚨 QUAN TRỌNG: ANH CHỊ NHỚ VÀO GROUP ZALO ĐỂ TIỆN THEO DÕI THÔNG BÁO NHA! 🚨<br>
                                <a href='{zalo_link}' target='_blank' style='color: #FFFFFF; text-decoration: none; font-size: 17px; display: inline-block; margin-top: 15px; background-color: #D32F2F; padding: 12px 25px; border-radius: 30px; box-shadow: 0 4px 6px rgba(0,0,0,0.2); transition: 0.3s;'>👉 BẤM VÀO ĐÂY ĐỂ THAM GIA GROUP 👈</a>
                            </div>
                            """, unsafe_allow_html=True)

                        st.markdown(f"""
                        <div style='text-align: center; margin: 25px 0 15px 0; padding: 10px; border-top: 2px dashed #2C5E1A; border-bottom: 2px dashed #2C5E1A;'>
                            <h4 style='color: #2C5E1A; margin: 0; font-size: 17px;'>🌿 BTC xin gửi lại thông tin cần thiết cho chuyến đi 🌿</h4>
                            <p style='color: #D4AF37; font-weight: bold; margin-top: 5px; font-size: 15px;'>Nhà mình nhớ lưu lại hình ảnh nha!</p>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        if thongtin_base64:
                            st.markdown(f"<img src='data:image/jpeg;base64,{thongtin_base64}' style='width:100%; border-radius:10px; box-shadow: 0 4px 8px rgba(0,0,0,0.1);'>", unsafe_allow_html=True)
        else:
            st.error("Không tìm thấy số điện thoại này. Vui lòng kiểm tra lại!")

# ================= PHẦN 2: ADMIN ONLY =================
st.sidebar.markdown("### ⚙️ FOR ADMIN ONLY")
admin_pass = st.sidebar.text_input("Nhập mật khẩu:", type="password")

def get_int(df, r, c):
    try:
        val = df.loc[r, c]
        return int(float(val)) if pd.notna(val) else 0
    except:
        return 0

if admin_pass == "0519":
    st.sidebar.success("Xác thực thành công!")
    st.markdown("---")
    st.markdown("## KHU VỰC QUẢN TRỊ ADMIN")
    
    admin_tabs = st.tabs(["📊 Thống kê Số Lượng", "💰 Thống kê Chi Phí", "📝 Chỉnh sửa Cấu hình gốc"])
    
    list_dots = df_config['Nội dung option'].apply(get_prefix).dropna().unique().tolist()
    list_dots = [d for d in list_dots if d.strip()]
    
    # TAB A: THỐNG KÊ SỐ LƯỢNG
    with admin_tabs[0]:
        selected_dot_A = st.selectbox("Chọn Đợt để xem thống kê:", ["All"] + list_dots, key="sel_A")
        
        def render_stats_for_dot(dot_prefix, df_full):
            st.markdown(f"<div class='section-title'>{dot_prefix.upper()}</div>", unsafe_allow_html=True)
            
            df_filtered = df_full[df_full[col_dot].apply(get_prefix).apply(normalize_text) == normalize_text(dot_prefix)]
            
            limit_mask = df_config['Nội dung option'].apply(get_prefix).apply(normalize_text) == normalize_text(dot_prefix)
            limit_series = df_config.loc[limit_mask, 'SL giới hạn']
            limit_val = int(float(limit_series.values[0])) if not limit_series.empty and pd.notna(limit_series.values[0]) else 0
            
            sl_dang_cap_nhat = df_filtered[df_filtered['Trạng thái CK'] == 'ĐANG CẬP NHẬT']['SL'].apply(pd.to_numeric, errors='coerce').sum()
            sl_chot_don = df_filtered[df_filtered['Trạng thái CK'] == 'CHỐT ĐƠN THÀNH CÔNG']['SL'].apply(pd.to_numeric, errors='coerce').sum()
            sl_huy = df_filtered[df_filtered['Trạng thái CK'] == 'HỦY SLOT']['SL'].apply(pd.to_numeric, errors='coerce').sum()
            sl_tong = sl_dang_cap_nhat + sl_chot_don
            
            dt_dukien = df_filtered[df_filtered['Trạng thái CK'].isin(['ĐANG CẬP NHẬT', 'CHỐT ĐƠN THÀNH CÔNG'])]['Tổng tiền'].apply(parse_money).sum()
            dt_thucte = df_filtered[df_filtered['Trạng thái CK'] == 'CHỐT ĐƠN THÀNH CÔNG']['Tổng tiền'].apply(parse_money).sum()
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"""
                <div class='metric-card'>
                    <div class='metric-title'>SỐ LƯỢNG ĐĂNG KÝ</div>
                    <div class='metric-value'>{int(sl_tong)} / {limit_val}</div>
                    <div style='font-size:14px; margin-top:10px;'>
                        🟢 Đang cập nhật: {int(sl_dang_cap_nhat)} <br>
                        ✅ Chốt đơn: {int(sl_chot_don)} <br>
                        ❌ Hủy slot: {int(sl_huy)}
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
            with col2:
                st.markdown(f"""
                <div class='metric-card'>
                    <div class='metric-title'>DOANH THU</div>
                    <div class='metric-value'>{dt_thucte:,.0f} đ</div>
                    <div style='font-size:14px; margin-top:10px;'>
                        Dự kiến thu: {dt_dukien:,.0f} đ
                    </div>
                </div>
                """, unsafe_allow_html=True)
            
            st.markdown("<h4 style='margin-top: 15px;'>Danh sách khách đã chốt đơn</h4>", unsafe_allow_html=True)
            df_chot = df_filtered[df_filtered['Trạng thái CK'] == 'CHỐT ĐƠN THÀNH CÔNG']
            
            if not df_chot.empty:
                ds_html = "<table class='custom-table' style='margin-top:0;'><thead><tr><th>Họ và Tên</th><th>Năm sinh</th></tr></thead><tbody>"
                for _, r in df_chot.iterrows():
                    for ng in str(r['DS người']).split(','):
                        parts = ng.split('-')
                        t = parts[0].strip() if len(parts)>0 else ""
                        ns = parts[1].strip() if len(parts)>1 else ""
                        ds_html += f"<tr><td>{t}</td><td>{ns}</td></tr>"
                ds_html += "</tbody></table>"
                st.markdown(ds_html, unsafe_allow_html=True)
                
                word_data = generate_word(df_chot, dot_prefix)
                st.download_button(
                    label=f"📥 TẢI XUỐNG DANH SÁCH (WORD) - {dot_prefix.upper()}",
                    data=word_data,
                    file_name=f"Danh_sach_chot_don_{dot_prefix}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            else:
                st.info("Chưa có khách hàng chốt đơn thành công.")
                
        if selected_dot_A == "All":
            for d in list_dots:
                render_stats_for_dot(d, df_data)
        else:
            render_stats_for_dot(selected_dot_A, df_data)

    # TAB B: THỐNG KÊ CHI PHÍ
    with admin_tabs[1]:
        selected_dot_B = st.selectbox("Chọn Đợt tính chi phí:", list_dots, key="sel_B")
        
        df_dot_B = df_data[(df_data[col_dot].apply(get_prefix).apply(normalize_text) == normalize_text(selected_dot_B)) & (df_data['Trạng thái CK'] == 'CHỐT ĐƠN THÀNH CÔNG')]
        khach_chot_don = df_dot_B['SL'].apply(pd.to_numeric, errors='coerce').sum()
        khach_chot_don = int(khach_chot_don) if pd.notna(khach_chot_don) else 0
        
        doanhthu_thucte_B = df_dot_B['Tổng tiền'].apply(parse_money).sum()
        
        st.write(f"**Tổng số người (CHỐT ĐƠN THÀNH CÔNG):** {khach_chot_don} người")
        
        df_cp = df_config[['Chi phí tổ chức', 'Unit cost', 'Loại chi phí']].dropna(subset=['Chi phí tổ chức'])
        df_cp = df_cp[df_cp['Chi phí tổ chức'] != '']
        
        list_dropdowns = df_cp[df_cp['Loại chi phí'] == 'Dropdown']['Chi phí tổ chức'].tolist()
        selected_opts = st.multiselect("Chọn các hạng mục phát sinh (Dropdown):", list_dropdowns)
        
        cost_data = []
        total_cost = 0
        
        for _, r in df_cp.iterrows():
            ten_cp = r['Chi phí tổ chức']
            u_cost = parse_money(r['Unit cost'])
            l_cp = r['Loại chi phí']
            
            if l_cp == "Cố định":
                sl = 1
                thanh_tien = u_cost * sl
                cost_data.append([ten_cp, f"{u_cost:,.0f}", sl, f"{thanh_tien:,.0f}"])
                total_cost += thanh_tien
            elif l_cp == "Theo đầu người":
                sl = khach_chot_don
                thanh_tien = u_cost * sl
                cost_data.append([ten_cp, f"{u_cost:,.0f}", sl, f"{thanh_tien:,.0f}"])
                total_cost += thanh_tien
            elif l_cp == "Dropdown" and ten_cp in selected_opts:
                sl = 1
                thanh_tien = u_cost * sl
                cost_data.append([ten_cp, f"{u_cost:,.0f}", sl, f"{thanh_tien:,.0f}"])
                total_cost += thanh_tien
                
        loi_nhuan = doanhthu_thucte_B - total_cost
        
        st.markdown("<div class='section-title'>BẢNG DỰ TOÁN CHI PHÍ</div>", unsafe_allow_html=True)
        
        table_html = "<table class='custom-table'><thead><tr><th>Chi phí</th><th>Unit Cost</th><th>Số lượng</th><th>Thành tiền (VNĐ)</th></tr></thead><tbody>"
        for row in cost_data:
            table_html += f"<tr><td>{row[0]}</td><td>{row[1]}</td><td>{row[2]}</td><td>{row[3]}</td></tr>"
        
        table_html += f"<tr class='total-row'><td>TỔNG CỘNG CHI PHÍ</td><td></td><td></td><td>{total_cost:,.0f}</td></tr>"
        table_html += f"<tr class='profit-row'><td>LỢI NHUẬN ƯỚC TÍNH</td><td></td><td></td><td>{loi_nhuan:,.0f}</td></tr>"
        table_html += "</tbody></table>"
        
        st.markdown(table_html, unsafe_allow_html=True)

    # TAB C: CHỈNH SỬA GỐC
    with admin_tabs[2]:
        st.warning("⚠️ Mọi thay đổi ở đây sẽ ghi đè trực tiếp lên tab 'Thông số' của Google Sheet.")
        df_update = df_config.copy()
        
        with st.form("update_config_form"):
            st.markdown("### 1. Cụm Scheme vé")
            col_v1, col_v2, col_v3 = st.columns(3)
            with col_v1: ve = st.number_input("Vé 1 người", value=get_int(df_update, 0, 'Vé 1 người'), step=1)
            with col_v2: sl_km = st.number_input("SL khuyến mãi", value=get_int(df_update, 0, 'SL khuyến mãi'), step=1)
            with col_v3: km = st.number_input("Scheme khuyến mãi", value=get_int(df_update, 0, 'Scheme khuyến mãi'), step=1)
            
            st.markdown("### 2. Cụm Thời gian")
            col_t1, col_t2 = st.columns(2)
            with col_t1: tg_ck = st.number_input("Thời gian CK (phút)", value=get_int(df_update, 0, 'Thời gian CK'), step=1)
            with col_t2: tg_check = st.number_input("Thời gian Check (phút)", value=get_int(df_update, 0, 'Thời gian Check'), step=1)
            
            st.markdown("### 3. Cụm Số lượng giới hạn & Link Zalo (2 Đợt)")
            col_d1, col_d2, col_d3 = st.columns([2, 1, 2])
            
            d1_name = str(df_update.loc[0, 'Nội dung option']) if len(df_update) > 0 and pd.notna(df_update.loc[0, 'Nội dung option']) else ""
            prefix1 = get_prefix(d1_name) if d1_name.strip() else "Đợt 1"
            d1_limit = get_int(df_update, 0, 'SL giới hạn')
            d1_zalo = str(df_update.loc[0, 'Link Zalo']) if len(df_update) > 0 and pd.notna(df_update.loc[0, 'Link Zalo']) else ""
            
            d2_name = str(df_update.loc[1, 'Nội dung option']) if len(df_update) > 1 and pd.notna(df_update.loc[1, 'Nội dung option']) else ""
            prefix2 = get_prefix(d2_name) if d2_name.strip() else "Đợt 2"
            d2_limit = get_int(df_update, 1, 'SL giới hạn')
            d2_zalo = str(df_update.loc[1, 'Link Zalo']) if len(df_update) > 1 and pd.notna(df_update.loc[1, 'Link Zalo']) else ""
            
            with col_d1: opt1 = st.text_input(f"Tên {prefix1}", value=str(d1_name))
            with col_d2: lim1 = st.number_input(f"SL {prefix1}", value=d1_limit, step=1)
            with col_d3: zalo1 = st.text_input(f"Link Zalo {prefix1}", value=str(d1_zalo))
            
            with col_d1: opt2 = st.text_input(f"Tên {prefix2}", value=str(d2_name))
            with col_d2: lim2 = st.number_input(f"SL {prefix2}", value=d2_limit, step=1)
            with col_d3: zalo2 = st.text_input(f"Link Zalo {prefix2}", value=str(d2_zalo))
            
            st.markdown("### 4. Cụm Chi phí tổ chức")
            st.caption("Chỉnh sửa trực tiếp trên bảng bên dưới (Có thể click đúp vào ô để gõ)")
            
            cp_cols = ['Chi phí tổ chức', 'Unit cost', 'Loại chi phí']
            if 'Loại chi phí' not in df_update.columns:
                df_update['Loại chi phí'] = ""
                
            df_cp_edit = df_update[cp_cols].copy().dropna(how='all')
            if len(df_cp_edit) < 10: 
                extra_rows = pd.DataFrame([["", 0, "Cố định"]]*5, columns=cp_cols)
                df_cp_edit = pd.concat([df_cp_edit, extra_rows], ignore_index=True)
                
            edited_cp = st.data_editor(
                df_cp_edit, 
                column_config={
                    "Loại chi phí": st.column_config.SelectboxColumn(
                        "Loại chi phí", options=["Cố định", "Theo đầu người", "Dropdown"], required=True
                    )
                },
                num_rows="dynamic",
                use_container_width=True
            )
            
            submitted = st.form_submit_button("LƯU THAY ĐỔI LÊN GOOGLE SHEET 💾")
            if submitted:
                df_update.loc[0, 'Vé 1 người'] = ve
                df_update.loc[0, 'SL khuyến mãi'] = sl_km
                df_update.loc[0, 'Scheme khuyến mãi'] = km
                
                df_update.loc[0, 'Thời gian CK'] = tg_ck
                df_update.loc[0, 'Thời gian Check'] = tg_check
                
                df_update.loc[0, 'Nội dung option'] = opt1
                df_update.loc[0, 'SL giới hạn'] = lim1
                df_update.loc[0, 'Link Zalo'] = zalo1
                
                df_update.loc[1, 'Nội dung option'] = opt2
                df_update.loc[1, 'SL giới hạn'] = lim2
                df_update.loc[1, 'Link Zalo'] = zalo2
                
                df_update['Chi phí tổ chức'] = None
                df_update['Unit cost'] = None
                df_update['Loại chi phí'] = None
                
                for i, r in edited_cp.iterrows():
                    if pd.notna(r['Chi phí tổ chức']) and str(r['Chi phí tổ chức']).strip() != "":
                        df_update.loc[i, 'Chi phí tổ chức'] = r['Chi phí tổ chức']
                        df_update.loc[i, 'Unit cost'] = r['Unit cost']
                        df_update.loc[i, 'Loại chi phí'] = r['Loại chi phí']
                
                try:
                    conn.update(worksheet="Thông số", data=df_update)
                    st.success("✅ Đã cập nhật dữ liệu thành công lên Google Sheet! Hãy làm mới trang để thấy thay đổi.")
                    st.cache_data.clear()
                except Exception as e:
                    st.error(f"Lỗi khi cập nhật lên Google Sheet: {e}")
